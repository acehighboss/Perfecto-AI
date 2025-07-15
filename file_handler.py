import os
import requests
import tempfile
from typing import List, Optional
from io import BytesIO
import streamlit as st

# LangChain imports
from langchain_core.documents import Document
from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders.base import BaseLoader

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

class WebLoader(BaseLoader):
    """웹 페이지 로더"""
    
    def __init__(self, url: str):
        self.url = url
    
    def load(self) -> List[Document]:
        """웹 페이지에서 문서 로드"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(self.url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # 불필요한 태그 제거
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            if text.strip():
                return [Document(
                    page_content=text,
                    metadata={"source": self.url, "type": "web"}
                )]
            else:
                return []
                
        except Exception as e:
            st.error(f"URL 처리 중 오류 발생: {str(e)}")
            return []

class StreamlitFileLoader(BaseLoader):
    """Streamlit 업로드 파일 로더"""
    
    def __init__(self, uploaded_file):
        self.uploaded_file = uploaded_file
        self.file_extension = uploaded_file.name.split('.')[-1].lower()
    
    def load(self) -> List[Document]:
        """업로드된 파일에서 문서 로드"""
        try:
            if self.file_extension == 'pdf':
                return self._load_pdf()
            elif self.file_extension in ['docx', 'doc']:
                return self._load_docx()
            elif self.file_extension == 'txt':
                return self._load_txt()
            else:
                st.error(f"지원하지 않는 파일 형식: {self.file_extension}")
                return []
        except Exception as e:
            st.error(f"파일 처리 중 오류 발생: {str(e)}")
            return []
    
    def _load_pdf(self) -> List[Document]:
        """PDF 파일 로드"""
        try:
            # 임시 파일로 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(self.uploaded_file.read())
                tmp_file.flush()
                
                # PyMuPDF로 직접 처리 (더 나은 제어를 위해)
                pdf_document = fitz.open(tmp_file.name)
                documents = []
                
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    page_text = page.get_text()
                    
                    if page_text.strip():
                        documents.append(Document(
                            page_content=page_text,
                            metadata={
                                "source": self.uploaded_file.name,
                                "page": page_num + 1,
                                "type": "pdf"
                            }
                        ))
                
                pdf_document.close()
                os.unlink(tmp_file.name)
                
                return documents
                
        except Exception as e:
            st.error(f"PDF 처리 중 오류: {str(e)}")
            return []
    
    def _load_docx(self) -> List[Document]:
        """Word 문서 로드"""
        try:
            doc_file = BytesIO(self.uploaded_file.read())
            doc = DocxDocument(doc_file)
            
            # 문단 텍스트 추출
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # 표 텍스트 추출
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # 전체 텍스트 결합
            full_text = "\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n" + "\n\n".join(tables_text)
            
            if full_text.strip():
                return [Document(
                    page_content=full_text,
                    metadata={
                        "source": self.uploaded_file.name,
                        "type": "docx"
                    }
                )]
            else:
                return []
                
        except Exception as e:
            st.error(f"Word 문서 처리 중 오류: {str(e)}")
            return []
    
    def _load_txt(self) -> List[Document]:
        """텍스트 파일 로드"""
        try:
            file_content = self.uploaded_file.read()
            
            # 다양한 인코딩 시도
            encodings = ['utf-8', 'cp949', 'latin-1']
            text = None
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                st.error("텍스트 파일의 인코딩을 인식할 수 없습니다.")
                return []
            
            if text.strip():
                return [Document(
                    page_content=text,
                    metadata={
                        "source": self.uploaded_file.name,
                        "type": "txt"
                    }
                )]
            else:
                return []
                
        except Exception as e:
            st.error(f"텍스트 파일 처리 중 오류: {str(e)}")
            return []

class FileHandler:
    def __init__(self):
        # 텍스트 스플리터 설정
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def load_url(self, url: str) -> List[Document]:
        """URL에서 문서 로드"""
        loader = WebLoader(url)
        documents = loader.load()
        return self.split_documents(documents)
    
    def load_file(self, uploaded_file) -> List[Document]:
        """업로드된 파일에서 문서 로드"""
        loader = StreamlitFileLoader(uploaded_file)
        documents = loader.load()
        return self.split_documents(documents)
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """문서를 청크로 분할"""
        if not documents:
            return []
        
        try:
            split_docs = self.text_splitter.split_documents(documents)
            
            # 메타데이터에 청크 정보 추가
            for i, doc in enumerate(split_docs):
                doc.metadata["chunk_id"] = i
                doc.metadata["chunk_size"] = len(doc.page_content)
            
            return split_docs
            
        except Exception as e:
            st.error(f"문서 분할 중 오류: {str(e)}")
            return []
    
    def get_document_info(self, documents: List[Document]) -> dict:
        """문서 정보 반환"""
        if not documents:
            return {"total_chunks": 0, "total_characters": 0, "sources": []}
        
        total_chars = sum(len(doc.page_content) for doc in documents)
        sources = list(set(doc.metadata.get("source", "Unknown") for doc in documents))
        
        return {
            "total_chunks": len(documents),
            "total_characters": total_chars,
            "sources": sources
        }
